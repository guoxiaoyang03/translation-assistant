"""Export module — save translation results in various formats.

Supported formats: .txt (UTF-8), .docx/.doc (Word), .pdf
"""

from pathlib import Path
import platform
from typing import Optional


# ── CJK font discovery ─────────────────────────────────────────────

def _find_cjk_font() -> Optional[str]:
    """Return the path to a CJK-capable font on the current system, or None."""
    system = platform.system()

    if system == "Windows":
        candidates = [
            "C:/Windows/Fonts/msyh.ttc",    # Microsoft YaHei
            "C:/Windows/Fonts/msyhbd.ttc",  # Microsoft YaHei Bold
            "C:/Windows/Fonts/simsun.ttc",  # SimSun
            "C:/Windows/Fonts/simhei.ttf",  # SimHei
        ]
    elif system == "Darwin":
        candidates = [
            "/System/Library/Fonts/PingFang.ttc",
            "/System/Library/Fonts/STHeiti Light.ttc",
            "/System/Library/Fonts/Hiragino Sans GB.ttc",
        ]
    else:  # Linux / others
        candidates = [
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/noto-cjk/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/wqy-microhei/wqy-microhei.ttc",
        ]

    for path in candidates:
        if Path(path).exists():
            return path
    return None


# ── Public API ─────────────────────────────────────────────────────

def export_translation(content: str, filepath: str) -> None:
    """Export *content* to *filepath*.

    The format is determined by the file extension:
    - ``.txt``  → UTF-8 plain text
    - ``.docx`` / ``.doc`` → Word document (via python-docx)
    - ``.pdf``  → PDF document (via fpdf2)
    """
    ext = Path(filepath).suffix.lower()

    if ext == ".txt":
        _save_txt(content, filepath)
    elif ext in (".docx", ".doc"):
        _save_docx(content, filepath)
    elif ext == ".pdf":
        _save_pdf(content, filepath)
    else:
        raise ValueError(f"不支持的文件格式: {ext}")


# ── Format-specific writers ────────────────────────────────────────

def _save_txt(content: str, filepath: str) -> None:
    """Write UTF-8 plain text."""
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(content)


def _save_docx(content: str, filepath: str) -> None:
    """Write a Word document, preserving line breaks."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "缺少 python-docx 库，请执行: pip install python-docx"
        )

    doc = Document()
    for line in content.split("\n"):
        if line.strip():
            doc.add_paragraph(line)
        else:
            doc.add_paragraph("")  # preserve blank lines
    doc.save(filepath)


def _save_pdf(content: str, filepath: str) -> None:
    """Write a PDF document with CJK font support."""
    try:
        from fpdf import FPDF
    except ImportError:
        raise ImportError(
            "缺少 fpdf2 库，请执行: pip install fpdf2"
        )

    pdf = FPDF()
    pdf.add_page()

    font_path = _find_cjk_font()
    if font_path:
        pdf.add_font("cjk", fname=font_path)
        pdf.set_font("cjk", size=12)
    else:
        # Fallback: Helvetica won't render CJK, but at least won't crash
        pdf.set_font("Helvetica", size=12)

    for line in content.split("\n"):
        # fpdf2 with CJK fonts does not always reset x after multi_cell;
        # explicitly reposition to the left margin before each line.
        pdf.set_x(pdf.l_margin)
        if line.strip():
            pdf.multi_cell(0, 10, line)
        else:
            pdf.ln(10)

    pdf.output(filepath)
