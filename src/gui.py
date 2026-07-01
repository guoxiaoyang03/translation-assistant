"""GUI module — provides a topmost overlay window for interactive translation.

Features
--------
* Translation tab  — screenshot OCR / paste-or-import text → translate → save
* History tab      — browse, search, delete, and reload past translations
* System tray      — keep the app running in the background
* Global hotkey    — ``Ctrl+Shift+T`` anywhere to translate selected text
"""

import tkinter as tk
import tkinter.ttk as ttk
import threading
import time
from datetime import datetime
from pathlib import Path
from tkinter import messagebox, filedialog

from src.screenshot import capture_region
from src.translator import translate_image, translate_text
from src.exporter import export_translation
from src.history import (
    init_db,
    add_translation,
    get_translations,
    count_translations,
    delete_translation,
)
from src.hotkey import HotkeyService

# ══════════════════════════════════════════════════════════════════
#  Color scheme
# ══════════════════════════════════════════════════════════════════
BG_MAIN        = "#F5F5F5"
BG_BTN         = "#4A90D9"
BG_BTN_HOVER   = "#357ABD"
BG_TAB_ACTIVE  = "#4A90D9"
BG_TAB_INACTIVE = "#E0E0E0"
FG_BTN         = "#FFFFFF"
FG_DISABLED    = "#B0B0B0"
FG_TAB_ACTIVE  = "#FFFFFF"
FG_TAB_INACTIVE = "#555555"
BG_TEXT        = "#FFFFFF"
FG_TEXT        = "#333333"
BORDER         = "#D0D0D0"

# ══════════════════════════════════════════════════════════════════
#  Fonts
# ══════════════════════════════════════════════════════════════════
FONT_NORMAL   = ("Microsoft YaHei", 10)
FONT_HEADING  = ("Microsoft YaHei", 12, "bold")
FONT_SMALL    = ("Microsoft YaHei", 9)


def create_gui():
    """Create and run the main GUI loop."""

    root = tk.Tk()
    root.title("翻译助手 — Ctrl+Shift+Y 划词 | Ctrl+Shift+S 截图")
    root.configure(bg=BG_MAIN)
    root.minsize(480, 500)

    init_db()

    _t0 = time.time()  # startup timer

    # ── State ───────────────────────────────────────────────
    _translation_text: str = ""

    # Cache OCR result from screenshot pre-fill so on_translate can
    # skip the OCR step when the source text hasn't been edited.
    _last_ocr_text: str = ""

    # Direction
    direction_var = tk.StringVar(value="en2zh")

    def _get_langs() -> tuple[str, str]:
        if direction_var.get() == "en2zh":
            return ("en", "zh-Hans")
        else:
            return ("zh-CN", "en")

    # Thread-safe direction cache for the hotkey callback
    _dir_cache = {"source": "en", "target": "zh-Hans", "label": "英→中"}

    def _sync_direction(*_args) -> None:
        src, tgt = _get_langs()
        _dir_cache["source"] = src
        _dir_cache["target"] = tgt
        _dir_cache["label"] = (
            "英→中" if direction_var.get() == "en2zh" else "中→英"
        )

    direction_var.trace_add("write", _sync_direction)

    # Topmost toggle
    topmost_var = tk.BooleanVar(value=False)

    def _on_toggle_topmost() -> None:
        root.attributes("-topmost", topmost_var.get())

    # ── Styled button factory ───────────────────────────────

    def _make_btn(parent: tk.Misc, text: str, command) -> tk.Button:
        return tk.Button(
            parent,
            text=text,
            font=FONT_NORMAL,
            bg=BG_BTN,
            fg=FG_BTN,
            activebackground=BG_BTN_HOVER,
            activeforeground=FG_BTN,
            disabledforeground=FG_DISABLED,
            relief="flat",
            cursor="hand2",
            padx=15,
            pady=6,
            command=command,
        )

    # ── Output helper ───────────────────────────────────────

    def set_output(text: str, *, is_translation: bool = False) -> None:
        nonlocal _translation_text
        text_output.configure(state=tk.NORMAL)
        text_output.delete("1.0", tk.END)
        text_output.insert("1.0", text)
        text_output.configure(state=tk.DISABLED)
        if is_translation:
            _translation_text = text

    def _add_history_async(src: str, trans: str,
                           sl: str, tl: str, d: str) -> None:
        """Write a translation record on a background thread."""
        threading.Thread(
            target=add_translation,
            args=(src, trans, sl, tl, d),
            daemon=True,
        ).start()

    # ── Hotkey popup ────────────────────────────────────────

    def _show_hotkey_popup(source: str, translated: str) -> None:
        popup = tk.Toplevel(root)
        popup.title("翻译助手 — 划词翻译")
        popup.attributes("-topmost", True)
        popup.configure(bg=BG_MAIN)
        popup.resizable(True, True)

        # ── Position near cursor ──────────────────────────────
        px, py = root.winfo_pointerxy()
        screen_h = popup.winfo_screenheight()
        max_h = int(screen_h * 0.5)
        popup_w = 450

        # ── Estimate content height ───────────────────────────
        def _est_lines(text: str, chars_per_line: int = 45) -> int:
            n = 0
            for para in text.split("\n"):
                if not para:
                    n += 1
                else:
                    n += max(1, -(-len(para) // chars_per_line))
            return n

        src_est = _est_lines(source) + 1
        out_est = _est_lines(translated) + 1
        src_h = min(max(src_est, 2), 8)   # 2–8 lines
        out_h = min(max(out_est, 3), 12)  # 3–12 lines

        # ── Container ─────────────────────────────────────────
        container = tk.Frame(popup, bg=BG_MAIN, padx=10, pady=6)
        container.pack(fill=tk.BOTH, expand=True)

        # ── Source section ────────────────────────────────────
        lbl_src = tk.Label(container, text="📖 原文", font=FONT_HEADING,
                           bg=BG_MAIN, fg=BG_BTN, anchor="w")
        lbl_src.pack(fill=tk.X)

        frame_src = tk.Frame(container, bg=BG_MAIN)
        frame_src.pack(fill=tk.BOTH, expand=False, pady=(2, 8))

        src_scroll = tk.Scrollbar(frame_src)
        src_scroll.pack(side=tk.RIGHT, fill=tk.Y)

        src_text = tk.Text(
            frame_src, height=src_h, font=FONT_NORMAL,
            wrap=tk.WORD, bg=BG_TEXT, fg=FG_TEXT,
            relief="solid", borderwidth=1,
            yscrollcommand=src_scroll.set,
        )
        src_text.insert("1.0", source)
        src_text.configure(state=tk.DISABLED)
        src_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        src_scroll.configure(command=src_text.yview)

        # ── Translated section ────────────────────────────────
        lbl_out = tk.Label(container, text="🌐 译文", font=FONT_HEADING,
                           bg=BG_MAIN, fg=BG_BTN, anchor="w")
        lbl_out.pack(fill=tk.X)

        frame_out = tk.Frame(container, bg=BG_MAIN)
        frame_out.pack(fill=tk.BOTH, expand=True, pady=(2, 4))

        out_scroll = tk.Scrollbar(frame_out)
        out_scroll.pack(side=tk.RIGHT, fill=tk.Y)

        out_text = tk.Text(
            frame_out, height=out_h, font=FONT_NORMAL,
            wrap=tk.WORD, bg=BG_TEXT, fg=FG_TEXT,
            relief="solid", borderwidth=1,
            yscrollcommand=out_scroll.set,
        )
        out_text.insert("1.0", translated)
        out_text.configure(state=tk.DISABLED)
        out_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        out_scroll.configure(command=out_text.yview)

        # ── Hint ──────────────────────────────────────────────
        lbl_hint = tk.Label(
            container,
            text="点击空白处关闭  |  8 秒后自动消失",
            font=FONT_SMALL, bg=BG_MAIN, fg="#999999",
        )
        lbl_hint.pack(pady=(4, 0))

        # ── Click-on-blank-area dismiss ───────────────────────
        def _dismiss(_event=None) -> None:
            popup.destroy()

        for w in (container, lbl_src, lbl_out, lbl_hint):
            w.bind("<Button-1>", _dismiss)
        # Do NOT bind on Text widgets — allows text selection & scroll

        # ── Geometry & auto-close ─────────────────────────────
        popup.geometry(f"{popup_w}x{max_h}+{px + 20}+{py + 20}")
        popup.minsize(400, 200)
        popup.after(8000, lambda: popup.destroy() if popup.winfo_exists() else None)
        popup.focus_force()

    # ── Hotkey translate callback (main-thread safe) ────────

    def _on_hotkey_translate(text: str) -> None:
        """Called on the main thread when the hotkey fires."""
        src = _dir_cache["source"]
        tgt = _dir_cache["target"]
        label = _dir_cache["label"]

        try:
            result = translate_text(text, src, tgt)
        except Exception as exc:
            result = f"翻译失败：{exc}"

        # Save to history
        _add_history_async(text, result, src, tgt, label)

        # Show popup
        _show_hotkey_popup(text, result)

    # ── Button callbacks ────────────────────────────────────

    def _has_real_source() -> bool:
        """Return True if the source text box contains user content."""
        raw = text_source.get("1.0", tk.END).strip()
        return bool(raw) and not raw.startswith("[系统]")

    def _update_translate_btn() -> None:
        """Enable the translate button only when input is available."""
        if _has_real_source() or Path("data/raw/screenshot.png").is_file():
            btn_start.configure(state=tk.NORMAL)
        else:
            btn_start.configure(state=tk.DISABLED)

    def _contains_chinese(text: str) -> bool:
        """Return True if *text* contains any CJK Unified Ideograph."""
        return any("一" <= ch <= "鿿" for ch in text)

    def _apply_ocr_result(ocr_text: str) -> None:
        """Fill the source text box with OCR result (main-thread safe)."""
        nonlocal _last_ocr_text
        text_source.delete("1.0", tk.END)
        text_source.insert("1.0", ocr_text)
        _last_ocr_text = ocr_text
        # Auto-detect language & switch direction
        if _contains_chinese(ocr_text):
            direction_var.set("zh2en")
        else:
            direction_var.set("en2zh")
        _update_translate_btn()

    def on_select_region() -> None:
        # Clear source text for a clean state on new screenshot
        text_source.delete("1.0", tk.END)

        success = capture_region(root)

        if success or Path("data/raw/screenshot.png").is_file():
            set_output('[系统] 区域选择完成，请点击"开始翻译"。')
            text_source.insert("1.0", "[系统] OCR 识别中...")

            # Run OCR on a background thread so the UI stays responsive
            def _do_ocr() -> None:
                try:
                    screenshot = Path("data/raw/screenshot.png")
                    source_lang, target_lang = _get_langs()
                    _, ocr_text = translate_image(
                        str(screenshot), source_lang, target_lang
                    )
                except Exception:
                    ocr_text = ""
                if ocr_text and ocr_text not in ("未识别到文字", ""):
                    root.after(0, _apply_ocr_result, ocr_text)
                else:
                    root.after(0, lambda: (
                        text_source.delete("1.0", tk.END),
                        _update_translate_btn(),
                    ))

            threading.Thread(target=_do_ocr, daemon=True).start()
            _update_translate_btn()
        else:
            set_output("[系统] 已取消选择。")
            _update_translate_btn()

    def on_import() -> None:
        filepath = filedialog.askopenfilename(
            parent=root,
            filetypes=[
                ("文本文件 (*.txt)", "*.txt"),
                ("Word 文档 (*.docx)", "*.docx"),
            ],
            title="导入文本文件",
        )
        if not filepath:
            return
        try:
            ext = Path(filepath).suffix.lower()
            if ext == ".txt":
                content = Path(filepath).read_text(encoding="utf-8")
            elif ext == ".docx":
                from docx import Document
                doc = Document(filepath)
                content = "\n".join(
                    p.text for p in doc.paragraphs if p.text.strip()
                )
                if not content:
                    content = "\n".join(p.text for p in doc.paragraphs)
            else:
                messagebox.showerror("格式错误", f"不支持的文件格式: {ext}")
                return
        except Exception as exc:
            messagebox.showerror("导入失败", f"文件读取失败：{exc}")
            return
        text_source.delete("1.0", tk.END)
        text_source.insert("1.0", content)
        set_output(f"[系统] 已导入文件: {filepath}")
        _update_translate_btn()

    def _disable_buttons() -> None:
        btn_start.configure(text="翻译中...", state=tk.DISABLED)
        btn_select.configure(state=tk.DISABLED)
        btn_import.configure(state=tk.DISABLED)
        btn_save.configure(state=tk.DISABLED)

    def _enable_buttons() -> None:
        btn_start.configure(text="▶️ 开始翻译", state=tk.NORMAL)
        btn_select.configure(state=tk.NORMAL)
        btn_import.configure(state=tk.NORMAL)
        btn_save.configure(state=tk.NORMAL)

    def on_translate() -> None:
        source_lang, target_lang = _get_langs()
        direction_label = (
            "英→中" if direction_var.get() == "en2zh" else "中→英"
        )

        # ── Direct text translation ──────────────────────
        if _has_real_source():
            source_content = text_source.get("1.0", tk.END).strip()
            set_output("[系统] 正在翻译中，请稍候...")
            _disable_buttons()
            root.update_idletasks()

            try:
                result = translate_text(source_content, source_lang, target_lang)
            except Exception as exc:
                set_output(f"翻译出错：{exc}")
                return
            finally:
                _enable_buttons()

            set_output(result, is_translation=True)
            _add_history_async(
                source_content, result,
                source_lang, target_lang, direction_label,
            )
            return

        # ── Screenshot → OCR → translate ──────────────────
        screenshot = Path("data/raw/screenshot.png")
        if not screenshot.is_file():
            messagebox.showwarning(
                "无法翻译", "请先截图或输入源文本后再开始翻译。"
            )
            return

        set_output("[系统] 正在翻译中，请稍候...")
        _disable_buttons()
        root.update_idletasks()

        try:
            # If the source text still matches the cached OCR from
            # on_select_region() pre-fill, reuse it — no double OCR.
            if (_last_ocr_text
                    and text_source.get("1.0", tk.END).strip() == _last_ocr_text):
                ocr_text = _last_ocr_text
                result = translate_text(ocr_text, source_lang, target_lang)
            else:
                result, ocr_text = translate_image(
                    str(screenshot), source_lang, target_lang
                )
                # Update the source box with fresh OCR result
                text_source.delete("1.0", tk.END)
                text_source.insert("1.0", ocr_text)
        except Exception as exc:
            set_output(f"翻译出错：{exc}")
            return
        finally:
            _enable_buttons()

        set_output(result, is_translation=True)

        _add_history_async(
            ocr_text, result,
            source_lang, target_lang, direction_label,
        )

    def on_save() -> None:
        if not _translation_text.strip():
            messagebox.showwarning(
                "无法保存",
                "没有可保存的翻译内容。\n请先截图并翻译，或输入源文本翻译。",
            )
            return

        default_name = f"翻译结果_{datetime.now().strftime('%Y%m%d_%H%M')}"
        default_dir = Path("data/processed")
        default_dir.mkdir(parents=True, exist_ok=True)

        filepath = filedialog.asksaveasfilename(
            parent=root,
            initialdir=str(default_dir),
            initialfile=default_name,
            defaultextension=".txt",
            filetypes=[
                ("文本文件 (*.txt)", "*.txt"),
                ("Word 文档 (*.docx)", "*.docx"),
                ("Word 文档 旧版 (*.doc)", "*.doc"),
                ("PDF 文件 (*.pdf)", "*.pdf"),
            ],
            title="保存翻译文本",
        )
        if not filepath:
            return

        set_output("[系统] 正在保存文件...")
        try:
            export_translation(_translation_text, filepath)
        except Exception as exc:
            messagebox.showerror("保存失败", f"文件保存失败：{exc}")
            set_output(f"[系统] 保存失败：{exc}")
            return
        set_output(f"[系统] 文本保存完毕，路径为：{filepath}")

    # ═════════════════════════════════════════════════════════
    #  History tab helpers
    # ═════════════════════════════════════════════════════════

    _history_offset = 0
    _history_limit = 50
    _history_keyword: str = ""

    def _refresh_history() -> None:
        """Reload the history treeview from the database."""
        nonlocal _history_offset
        for item in tree_history.get_children():
            tree_history.delete(item)

        rows = get_translations(
            limit=_history_limit,
            offset=_history_offset,
            keyword=_history_keyword if _history_keyword else None,
        )
        for row in rows:
            rid, src, trans, direction, created = row
            # Truncate for display
            src_disp = src if len(src) <= 40 else src[:40] + "…"
            trans_disp = trans if len(trans) <= 40 else trans[:40] + "…"
            tree_history.insert(
                "", tk.END,
                iid=str(rid),
                values=(src_disp, trans_disp, direction, created),
            )

        total = count_translations(
            keyword=_history_keyword if _history_keyword else None
        )
        label_count.configure(text=f"共 {total} 条记录")
        has_more = _history_offset + _history_limit < total
        btn_more.configure(
            state=tk.NORMAL if has_more else tk.DISABLED
        )

    def _on_history_search() -> None:
        nonlocal _history_offset, _history_keyword
        _history_offset = 0
        _history_keyword = entry_search.get().strip()
        _refresh_history()

    def _on_history_next_page() -> None:
        nonlocal _history_offset
        _history_offset += _history_limit
        _refresh_history()

    def _on_history_delete() -> None:
        selected = tree_history.selection()
        if not selected:
            messagebox.showinfo("提示", "请先在列表中选中一条记录。")
            return
        if not messagebox.askyesno("确认删除", "确定要删除选中的翻译记录吗？"):
            return
        for iid in selected:
            delete_translation(int(iid))
        _refresh_history()

    def _on_history_double_click(_event=None) -> None:
        selected = tree_history.selection()
        if not selected:
            return
        # Get the full text from DB
        rows = get_translations(limit=1, offset=0)
        # We need to query by id — let's use a direct SQL approach
        import sqlite3
        rid = int(selected[0])
        conn = sqlite3.connect("data/translations.db")
        row = conn.execute(
            "SELECT source_text, translated_text FROM translations WHERE id = ?",
            (rid,),
        ).fetchone()
        conn.close()

        if row:
            text_source.delete("1.0", tk.END)
            text_source.insert("1.0", row[0])
            set_output(row[1], is_translation=True)
            _switch_tab("translate")

    # ═════════════════════════════════════════════════════════
    #  Tab switching
    # ═════════════════════════════════════════════════════════

    def _switch_tab(tab: str) -> None:
        if tab == "translate":
            btn_tab_translate.configure(bg=BG_TAB_ACTIVE, fg=FG_TAB_ACTIVE)
            btn_tab_history.configure(bg=BG_TAB_INACTIVE, fg=FG_TAB_INACTIVE)
            frame_history.pack_forget()
            frame_translate.pack(fill=tk.BOTH, expand=True)
        else:
            btn_tab_translate.configure(bg=BG_TAB_INACTIVE, fg=FG_TAB_INACTIVE)
            btn_tab_history.configure(bg=BG_TAB_ACTIVE, fg=FG_TAB_ACTIVE)
            frame_translate.pack_forget()
            frame_history.pack(fill=tk.BOTH, expand=True)
            _refresh_history()

    # ═════════════════════════════════════════════════════════
    #  Widgets — Top bar
    # ═════════════════════════════════════════════════════════

    # ── Topmost toggle ──────────────────────────────────
    cb_topmost = tk.Checkbutton(
        root,
        text="📌 窗口置顶",
        variable=topmost_var,
        command=_on_toggle_topmost,
        bg=BG_MAIN,
        activebackground=BG_MAIN,
        selectcolor=BG_MAIN,
        font=FONT_NORMAL,
        cursor="hand2",
    )
    cb_topmost.pack(pady=(8, 2), padx=16, anchor="w")

    # ── Tab bar ─────────────────────────────────────────
    frame_tab_bar = tk.Frame(root, bg=BG_MAIN)
    frame_tab_bar.pack(pady=(4, 0), padx=16, fill=tk.X)

    btn_tab_translate = tk.Button(
        frame_tab_bar,
        text="📝 翻译",
        font=FONT_NORMAL,
        bg=BG_TAB_ACTIVE,
        fg=FG_TAB_ACTIVE,
        activebackground=BG_TAB_ACTIVE,
        activeforeground=FG_TAB_ACTIVE,
        relief="flat",
        cursor="hand2",
        padx=20,
        pady=5,
        command=lambda: _switch_tab("translate"),
    )
    btn_tab_translate.pack(side=tk.LEFT)

    btn_tab_history = tk.Button(
        frame_tab_bar,
        text="📋 历史记录",
        font=FONT_NORMAL,
        bg=BG_TAB_INACTIVE,
        fg=FG_TAB_INACTIVE,
        activebackground=BG_TAB_ACTIVE,
        activeforeground=FG_TAB_ACTIVE,
        relief="flat",
        cursor="hand2",
        padx=20,
        pady=5,
        command=lambda: _switch_tab("history"),
    )
    btn_tab_history.pack(side=tk.LEFT)

    # ═════════════════════════════════════════════════════════
    #  TRANSLATE TAB
    # ═════════════════════════════════════════════════════════

    frame_translate = tk.Frame(root, bg=BG_MAIN)

    # ── Direction ───────────────────────────────────────
    frame_direction = tk.LabelFrame(
        frame_translate,
        text="翻译方向",
        font=FONT_HEADING,
        bg=BG_MAIN,
        fg=FG_TEXT,
    )
    frame_direction.pack(pady=(6, 6), padx=16, fill=tk.X)

    rb_en2zh = tk.Radiobutton(
        frame_direction,
        text="英 → 中  (English → 中文)",
        variable=direction_var,
        value="en2zh",
        font=FONT_NORMAL,
        bg=BG_MAIN,
        activebackground=BG_MAIN,
        selectcolor=BG_MAIN,
        anchor="w",
        cursor="hand2",
    )
    rb_en2zh.pack(side=tk.LEFT, padx=(10, 24), pady=6)

    rb_zh2en = tk.Radiobutton(
        frame_direction,
        text="中 → 英  (中文 → English)",
        variable=direction_var,
        value="zh2en",
        font=FONT_NORMAL,
        bg=BG_MAIN,
        activebackground=BG_MAIN,
        selectcolor=BG_MAIN,
        anchor="w",
        cursor="hand2",
    )
    rb_zh2en.pack(side=tk.LEFT, padx=(0, 10), pady=6)

    # ── Buttons ─────────────────────────────────────────
    btn_select = _make_btn(frame_translate, "📷 选择翻译区块", on_select_region)
    btn_select.pack(pady=(2, 6), padx=16, fill=tk.X)

    btn_import = _make_btn(frame_translate, "📂 导入文本", on_import)
    btn_import.pack(pady=6, padx=16, fill=tk.X)

    btn_start = _make_btn(frame_translate, "▶️ 开始翻译", on_translate)
    btn_start.configure(state=tk.DISABLED)   # disabled until input is available
    btn_start.pack(pady=6, padx=16, fill=tk.X)

    btn_save = _make_btn(frame_translate, "💾 保存文本", on_save)
    btn_save.pack(pady=6, padx=16, fill=tk.X)

    btn_exit = _make_btn(frame_translate, "❌ 退出关闭", None)
    btn_exit.pack(pady=6, padx=16, fill=tk.X)

    # ── Source text ─────────────────────────────────────
    frame_source = tk.LabelFrame(
        frame_translate,
        text="源文本",
        font=FONT_HEADING,
        bg=BG_MAIN,
        fg=FG_TEXT,
    )
    frame_source.pack(pady=(6, 6), padx=16, fill=tk.BOTH, expand=False)

    scrollbar_src = tk.Scrollbar(frame_source)
    scrollbar_src.pack(side=tk.RIGHT, fill=tk.Y)

    text_source = tk.Text(
        frame_source,
        height=5,
        font=FONT_NORMAL,
        wrap=tk.WORD,
        bg=BG_TEXT,
        fg=FG_TEXT,
        relief="solid",
        borderwidth=1,
        yscrollcommand=scrollbar_src.set,
    )
    text_source.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
    scrollbar_src.configure(command=text_source.yview)

    # Re-evaluate the translate button each time the user types
    text_source.bind("<KeyRelease>", lambda _e: _update_translate_btn())

    # ── Output ──────────────────────────────────────────
    frame_output = tk.LabelFrame(
        frame_translate,
        text="译文结果",
        font=FONT_HEADING,
        bg=BG_MAIN,
        fg=FG_TEXT,
    )
    frame_output.pack(pady=(6, 8), padx=16, fill=tk.BOTH, expand=True)

    scrollbar_out = tk.Scrollbar(frame_output)
    scrollbar_out.pack(side=tk.RIGHT, fill=tk.Y)

    text_output = tk.Text(
        frame_output,
        height=10,
        font=FONT_NORMAL,
        wrap=tk.WORD,
        state=tk.DISABLED,
        bg=BG_TEXT,
        fg=FG_TEXT,
        relief="solid",
        borderwidth=1,
        yscrollcommand=scrollbar_out.set,
    )
    text_output.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
    scrollbar_out.configure(command=text_output.yview)

    # Show translate tab by default
    frame_translate.pack(fill=tk.BOTH, expand=True)

    # ═════════════════════════════════════════════════════════
    #  HISTORY TAB
    # ═════════════════════════════════════════════════════════

    frame_history = tk.Frame(root, bg=BG_MAIN)

    # Search bar
    frame_search = tk.Frame(frame_history, bg=BG_MAIN)
    frame_search.pack(pady=(8, 4), padx=16, fill=tk.X)

    entry_search = tk.Entry(
        frame_search,
        font=FONT_NORMAL,
        bg=BG_TEXT,
        fg=FG_TEXT,
        relief="solid",
        borderwidth=1,
    )
    entry_search.pack(side=tk.LEFT, fill=tk.X, expand=True, ipady=3)
    entry_search.bind("<Return>", lambda e: _on_history_search())

    btn_search = tk.Button(
        frame_search,
        text="🔍 搜索",
        font=FONT_NORMAL,
        bg=BG_BTN,
        fg=FG_BTN,
        activebackground=BG_BTN_HOVER,
        activeforeground=FG_BTN,
        relief="flat",
        cursor="hand2",
        padx=14,
        pady=4,
        command=_on_history_search,
    )
    btn_search.pack(side=tk.LEFT, padx=(6, 0))

    # Treeview
    frame_tree = tk.Frame(frame_history, bg=BG_MAIN)
    frame_tree.pack(pady=4, padx=16, fill=tk.BOTH, expand=True)

    # ttk style for Treeview
    style = ttk.Style()
    style.configure(
        "History.Treeview",
        background=BG_TEXT,
        foreground=FG_TEXT,
        fieldbackground=BG_TEXT,
        font=FONT_NORMAL,
        rowheight=30,
    )
    style.configure(
        "History.Treeview.Heading",
        font=FONT_HEADING,
        background=BG_MAIN,
    )
    style.map("History.Treeview", background=[("selected", BG_BTN)])

    columns = ("source", "translated", "direction", "time")
    tree_history = ttk.Treeview(
        frame_tree,
        columns=columns,
        show="headings",
        selectmode="browse",
        style="History.Treeview",
    )
    tree_history.heading("source", text="原文")
    tree_history.heading("translated", text="译文")
    tree_history.heading("direction", text="方向")
    tree_history.heading("time", text="时间")

    tree_history.column("source", width=170, minwidth=100)
    tree_history.column("translated", width=170, minwidth=100)
    tree_history.column("direction", width=60, minwidth=50, anchor="center")
    tree_history.column("time", width=130, minwidth=100)

    tree_history.bind("<Double-1>", _on_history_double_click)

    scrollbar_hist = tk.Scrollbar(
        frame_tree, orient=tk.VERTICAL, command=tree_history.yview
    )
    tree_history.configure(yscrollcommand=scrollbar_hist.set)
    tree_history.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
    scrollbar_hist.pack(side=tk.RIGHT, fill=tk.Y)

    # Bottom bar
    frame_hist_bottom = tk.Frame(frame_history, bg=BG_MAIN)
    frame_hist_bottom.pack(pady=(4, 8), padx=16, fill=tk.X)

    btn_delete_hist = tk.Button(
        frame_hist_bottom,
        text="🗑 删除选中",
        font=FONT_NORMAL,
        bg="#E74C3C",
        fg=FG_BTN,
        activebackground="#C0392B",
        activeforeground=FG_BTN,
        relief="flat",
        cursor="hand2",
        padx=14,
        pady=4,
        command=_on_history_delete,
    )
    btn_delete_hist.pack(side=tk.LEFT)

    label_count = tk.Label(
        frame_hist_bottom,
        text="共 0 条记录",
        font=FONT_SMALL,
        bg=BG_MAIN,
        fg="#999999",
    )
    label_count.pack(side=tk.LEFT, padx=(12, 0))

    btn_more = tk.Button(
        frame_hist_bottom,
        text="加载更多 ▼",
        font=FONT_SMALL,
        bg=BG_MAIN,
        fg=BG_BTN,
        activebackground=BG_MAIN,
        activeforeground=BG_BTN_HOVER,
        relief="flat",
        cursor="hand2",
        padx=10,
        pady=2,
        state=tk.DISABLED,
        command=_on_history_next_page,
    )
    btn_more.pack(side=tk.RIGHT)

    # ═════════════════════════════════════════════════════════
    #  Window lifecycle  (defined before HotkeyService so
    #  _on_full_exit can be passed as the on_exit callback)
    # ═════════════════════════════════════════════════════════

    def _on_full_exit() -> None:
        """Full shutdown: stop hotkey + tray, then destroy the window."""
        hotkey_service.shutdown()
        root.destroy()

    def _on_window_close() -> None:
        """Hide to tray instead of quitting."""
        root.withdraw()

    root.protocol("WM_DELETE_WINDOW", _on_window_close)

    # Wire up the "退出关闭" button
    btn_exit.configure(command=_on_full_exit)

    # ═════════════════════════════════════════════════════════
    #  Hotkey service  (created last so all widgets exist)
    # ═════════════════════════════════════════════════════════

    hotkey_service = HotkeyService(
        root=root,
        on_translate=_on_hotkey_translate,
        on_screenshot=on_select_region,
        on_show_window=lambda: (root.deiconify(), root.lift()),
        on_exit=_on_full_exit,
    )
    hotkey_service.start()

    print(f"[启动] GUI 就绪，耗时 {time.time() - _t0:.2f}s，已进入主事件循环")
    root.mainloop()
